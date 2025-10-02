import streamlit as st
import pandas as pd
import re
import io
from typing import List, Dict, Tuple
import docx

# =============================================================================
# ФУНКЦИИ ЧТЕНИЯ ФАЙЛОВ
# =============================================================================

def read_uploaded_file(file):
    """
    Читает содержимое загруженного файла в зависимости от его типа
    """
    try:
        # Всегда сбрасываем позицию чтения файла
        file.seek(0)
        
        if file.name.endswith('.docx'):
            # Читаем DOCX файл
            doc = docx.Document(file)
            full_text = []
            
            # Читаем все параграфы
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Игнорируем пустые строки
                    full_text.append(paragraph.text)
            
            # Читаем таблицы, если есть
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            content = "\n".join(full_text)
            return content
        else:
            # Читаем TXT файл
            return file.getvalue().decode("utf-8")
            
    except Exception as e:
        st.error(f"Ошибка чтения файла {file.name}: {str(e)}")
        return None

# =============================================================================
# ФУНКЦИИ ОБРАБОТКИ ДАННЫХ
# =============================================================================

def parse_correct_order(file_content: str) -> List[Dict]:
    """
    Парсит правильный порядок образцов из файла с вырезками
    """
    if not file_content:
        return []
        
    lines = file_content.split('\n')
    correct_samples = []
    
    for line in lines:
        line = line.strip()
        
        # Пропускаем разделительные строки
        if re.match(r'^[-]+$', line) or not line:
            continue
            
        # Убираем разметку типа [ ]{.mark} - сначала убираем всю разметку
        clean_line = re.sub(r'\[(.*?)\]\{\.mark\}', r'\1', line)
        
        # Ищем строки с номерами и названиями образцов
        # Паттерн для строк типа: "13       [КПП ВД(50,А)]{.mark}"
        match = re.match(r'^\s*(\d+)\s+(.+)$', clean_line)
        if match:
            sample_number = int(match.group(1))
            sample_name = match.group(2).strip()
            
            correct_samples.append({
                'order': sample_number,
                'correct_name': sample_name,
                'key': create_sample_key(sample_name)
            })
    
    return correct_samples

def create_sample_key(sample_name: str) -> str:
    """
    Создает ключ для сопоставления образцов из разных источников
    """
    # Нормализуем название
    normalized = re.sub(r'\s+', ' ', sample_name.strip()).lower()
    
    # Извлекаем ключевые компоненты для сопоставления
    patterns = [
        r'([а-я]+)\s*([а-я]+)?\s*\((\d+[-\d]*),\s*([а-я])\)',  # КПП ВД(50,А)
        r'([а-я]+)\s*([а-я]+)?-?(\d+)?\s*\((\d+),\s*([а-я])\)', # КПП НД-1(19,А)
        r'(\d+)[,_]\s*([а-я]+)',  # Для формата "28_КПП ВД"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, normalized)
        if match:
            parts = [p for p in match.groups() if p]
            return '_'.join(parts)
    
    # Если паттерн не найден, ищем числа в названии
    numbers = re.findall(r'\d+', normalized)
    if numbers:
        return f"sample_{numbers[-1]}"
    
    return normalized

def parse_chemical_tables(file_content: str) -> Dict[str, List[Dict]]:
    """
    Парсит все таблицы химического анализа из файла
    Возвращает словарь: {марка_стали: [список_образцов]}
    """
    if not file_content:
        return {}
        
    lines = file_content.split('\n')
    tables = {}
    current_steel_grade = None
    current_table = []
    in_table = False
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Ищем марку стали
        steel_match = re.search(r'Марка стали:\s*([^\n]+)', line)
        if steel_match:
            if current_steel_grade and current_table:
                tables[current_steel_grade] = current_table
                current_table = []
            current_steel_grade = steel_match.group(1).strip()
            in_table = False
            continue
        
        # Ищем начало таблицы (строки с множеством дефисов)
        if re.match(r'^[-\\s]{10,}', line) and current_steel_grade:
            if not in_table:
                in_table = True
            continue
        
        # Парсим строки с образцами
        if in_table and current_steel_grade:
            # Пропускаем строки с требованиями ТУ
            if 'Требования ТУ' in line or '14-3Р-55-2001' in line:
                continue
                
            # Ищем строки с образцами (содержат номер и название)
            if re.match(r'^\s*\d+\s+[а-яА-Я]', line):
                # Разделяем по множественным пробелам
                parts = re.split(r'\s{2,}', line.strip())
                if len(parts) >= 3:
                    sample_data = {
                        'original_name': parts[1],
                        'measurements': parts[2:],
                        'key': create_sample_key(parts[1])
                    }
                    current_table.append(sample_data)
    
    # Добавляем последнюю таблицу
    if current_steel_grade and current_table:
        tables[current_steel_grade] = current_table
    
    return tables

def match_and_sort_samples(original_samples: List[Dict], correct_samples: List[Dict]) -> Tuple[List[Dict], List[Dict], List[Dict]]:
    """
    Сопоставляет и сортирует образцы по правильному порядку
    Возвращает: (сопоставленные, несопоставленные_правильные, несопоставленные_анализ)
    """
    key_to_correct = {}
    for correct in correct_samples:
        key_to_correct[correct['key']] = {
            'correct_name': correct['correct_name'],
            'order': correct['order']
        }
    
    matched_samples = []
    used_keys = set()
    
    for original in original_samples:
        if original['key'] in key_to_correct:
            matched_samples.append({
                'correct_name': key_to_correct[original['key']]['correct_name'],
                'measurements': original['measurements'],
                'order': key_to_correct[original['key']]['order'],
                'original_name': original['original_name'],
                'key': original['key']
            })
            used_keys.add(original['key'])
    
    # Сортируем по порядку из правильного списка
    matched_samples.sort(key=lambda x: x['order'])
    
    # Находим несопоставленные образцы из правильного порядка
    unmatched_correct = [correct for correct in correct_samples if correct['key'] not in used_keys]
    
    # Находим несопоставленные образцы из анализа
    unmatched_analysis = [original for original in original_samples if original['key'] not in key_to_correct]
    
    return matched_samples, unmatched_correct, unmatched_analysis

def create_matching_table(matched_samples: List[Dict]) -> pd.DataFrame:
    """
    Создает таблицу сопоставления названий
    """
    data = []
    for sample in matched_samples:
        data.append({
            '№ п/п': sample['order'],
            'Правильное название': sample['correct_name'],
            'Название из анализа': sample['original_name'],
            'Ключ сопоставления': sample['key']
        })
    
    return pd.DataFrame(data)

# =============================================================================
# STREAMLIT ИНТЕРФЕЙС
# =============================================================================

st.set_page_config(
    page_title="Сортировка химического анализа",
    page_icon="🔬",
    layout="wide"
)

st.markdown("""
<style>
    .main-header { 
        font-size: 2.5rem; 
        color: #1f77b4; 
        text-align: center; 
        margin-bottom: 2rem;
        font-weight: bold;
    }
    .stats-card { 
        background-color: #f0f2f6; 
        padding: 1rem; 
        border-radius: 0.5rem; 
        margin: 0.5rem 0;
        text-align: center;
    }
    .success-text { 
        color: #28a745; 
        font-weight: bold; 
        font-size: 1.2rem;
    }
    .warning-text { 
        color: #ffc107; 
        font-weight: bold;
        font-size: 1.2rem;
    }
    .file-info {
        background-color: #e8f4fd;
        padding: 0.5rem;
        border-radius: 0.3rem;
        margin: 0.3rem 0;
    }
    .debug-info {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
        font-family: monospace;
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)

def main():
    st.markdown('<div class="main-header">🔬 Автоматическая сортировка химического анализа</div>', unsafe_allow_html=True)
    
    # Описание приложения
    with st.expander("📖 ИНСТРУКЦИЯ ПО ИСПОЛЬЗОВАНИЮ", expanded=False):
        st.markdown("""
        ### Как использовать:
        1. **Загрузите файл с правильным порядком образцов** (TXT или DOCX формат)
        2. **Загрузите файл с результатами химического анализа** (TXT или DOCX формат)  
        3. **Нажмите кнопку "Обработать данные"**
        4. **Просмотрите таблицу сопоставления** - проверьте, правильно ли программа сопоставила образцы
        5. **Скачайте отсортированные таблицы**

        ### Принцип сопоставления:
        Программа ищет совпадения по **цифрам и буквам** в названиях образцов:
        - "КПП ВД 2, труба 28" → "КПП ВД(28,Г)" (сопоставление по числу 28)
        - "НГ 28_КПП ВД" → "КПП ВД(28,Г)" (сопоставление по числу 28)
        - "КПП ВД 2, труба 122" → "КПП ВД(50,А)" (сопоставление по числу 50)
        """)
    
    # Загрузка файлов
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📋 Файл с правильным порядком")
        correct_order_file = st.file_uploader(
            "Загрузите файл с правильным порядком образцов",
            type=['txt', 'docx'],
            key="correct_order",
            help="Файл должен содержать список образцов в правильном порядке с номерами"
        )
        
        if correct_order_file:
            st.markdown(f'<div class="file-info">✅ Файл загружен: <b>{correct_order_file.name}</b> ({correct_order_file.type})</div>', unsafe_allow_html=True)
    
    with col2:
        st.subheader("🧪 Файл с химическим анализом")
        chemical_analysis_file = st.file_uploader(
            "Загрузите файл с результатами химического анализа", 
            type=['txt', 'docx'],
            key="chemical_analysis",
            help="Файл должен содержать таблицы химического анализа по маркам стали"
        )
        
        if chemical_analysis_file:
            st.markdown(f'<div class="file-info">✅ Файл загружен: <b>{chemical_analysis_file.name}</b> ({chemical_analysis_file.type})</div>', unsafe_allow_html=True)
    
    # Кнопка обработки
    if st.button("🚀 ОБРАБОТАТЬ ДАННЫЕ", type="primary", use_container_width=True):
        if not correct_order_file or not chemical_analysis_file:
            st.error("❌ Пожалуйста, загрузите оба файла перед обработкой")
            return
        
        try:
            # Чтение файлов
            with st.spinner("📖 Читаю файлы..."):
                correct_order_content = read_uploaded_file(correct_order_file)
                chemical_analysis_content = read_uploaded_file(chemical_analysis_file)
            
            if not correct_order_content:
                st.error("❌ Не удалось прочитать файл с правильным порядком")
                return
                
            if not chemical_analysis_content:
                st.error("❌ Не удалось прочитать файл с химическим анализом")
                return
            
            # Парсинг
            with st.spinner("🔍 Анализирую данные..."):
                correct_samples = parse_correct_order(correct_order_content)
                chemical_tables = parse_chemical_tables(chemical_analysis_content)
            
            if not correct_samples:
                st.error("❌ Не удалось найти образцы в файле с правильным порядком")
                return
                
            if not chemical_tables:
                st.error("❌ Не удалось найти таблицы химического анализа")
                return
            
            # Обработка и сопоставление
            with st.spinner("🔄 Сопоставляю образцы..."):
                all_matched_samples = []
                all_unmatched_correct = []
                all_unmatched_analysis = []
                
                for steel_grade, samples in chemical_tables.items():
                    matched_samples, unmatched_correct, unmatched_analysis = match_and_sort_samples(samples, correct_samples)
                    all_matched_samples.extend(matched_samples)
                    all_unmatched_correct.extend(unmatched_correct)
                    all_unmatched_analysis.extend(unmatched_analysis)
            
            # СОЗДАЕМ ТАБЛИЦУ СОПОСТАВЛЕНИЯ
            st.subheader("🔍 ТАБЛИЦА СОПОСТАВЛЕНИЯ НАЗВАНИЙ")
            
            if all_matched_samples:
                # Создаем таблицу сопоставления
                matching_df = create_matching_table(all_matched_samples)
                
                # Показываем таблицу
                st.dataframe(matching_df, use_container_width=True)
                
                # Статистика сопоставления
                total_correct = len(correct_samples)
                total_matched = len(all_matched_samples)
                total_unmatched_correct = len(all_unmatched_correct)
                total_unmatched_analysis = len(all_unmatched_analysis)
                matching_rate = (total_matched / total_correct) * 100 if total_correct > 0 else 0
                
                st.success(f"✅ Сопоставлено {total_matched} из {total_correct} образцов ({matching_rate:.1f}%)")
                
                # Показываем несопоставленные образцы
                if all_unmatched_correct:
                    st.warning(f"⚠️ Не сопоставлено {total_unmatched_correct} образцов из правильного порядка:")
                    unmatched_data = []
                    for sample in all_unmatched_correct:
                        unmatched_data.append({
                            '№ п/п': sample['order'],
                            'Правильное название': sample['correct_name'],
                            'Ключ': sample['key']
                        })
                    st.dataframe(pd.DataFrame(unmatched_data), use_container_width=True)
                
                if all_unmatched_analysis:
                    st.warning(f"⚠️ Не сопоставлено {total_unmatched_analysis} образцов из анализа:")
                    unmatched_analysis_data = []
                    for sample in all_unmatched_analysis:
                        unmatched_analysis_data.append({
                            'Название из анализа': sample['original_name'],
                            'Ключ': sample['key']
                        })
                    st.dataframe(pd.DataFrame(unmatched_analysis_data), use_container_width=True)
                
                # Кнопка для продолжения - показать отсортированные таблицы
                if st.button("📊 ПОКАЗАТЬ ОТСОРТИРОВАННЫЕ ТАБЛИЦЫ", type="secondary"):
                    # Создаем финальные таблицы
                    final_tables = {}
                    for steel_grade, samples in chemical_tables.items():
                        matched_samples, _, _ = match_and_sort_samples(samples, correct_samples)
                        if matched_samples:
                            num_measurements = len(matched_samples[0]['measurements'])
                            columns = ['№', 'Образец'] + [f'Измерение {i+1}' for i in range(num_measurements)]
                            
                            data = []
                            for i, sample in enumerate(matched_samples):
                                row = [i+1, sample['correct_name']] + sample['measurements']
                                data.append(row)
                            
                            final_tables[steel_grade] = pd.DataFrame(data, columns=columns)
                    
                    # Показываем отсортированные таблицы
                    st.subheader("📋 ОТСОРТИРОВАННЫЕ РЕЗУЛЬТАТЫ")
                    
                    for steel_grade, table in final_tables.items():
                        with st.expander(f"🔩 Марка стали: {steel_grade} ({len(table)} образцов)", expanded=True):
                            st.dataframe(table, use_container_width=True)
                            
                            # Скачивание отдельной таблицы
                            excel_buffer = io.BytesIO()
                            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                                table.to_excel(writer, index=False, sheet_name=steel_grade[:30])
                            excel_buffer.seek(0)
                            
                            st.download_button(
                                label=f"📥 Скачать таблицу {steel_grade} (Excel)",
                                data=excel_buffer,
                                file_name=f"отсортировано_{steel_grade}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key=f"download_{steel_grade}"
                            )
                    
                    # Пакетное скачивание
                    if len(final_tables) > 1:
                        st.subheader("💾 ПАКЕТНОЕ СКАЧИВАНИЕ")
                        excel_buffer_all = io.BytesIO()
                        with pd.ExcelWriter(excel_buffer_all, engine='openpyxl') as writer:
                            for steel_grade, table in final_tables.items():
                                table.to_excel(writer, index=False, sheet_name=steel_grade[:30])
                        excel_buffer_all.seek(0)
                        
                        st.download_button(
                            label="📦 Скачать все таблицы (Excel)",
                            data=excel_buffer_all,
                            file_name="все_отсортированные_таблицы.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="download_all",
                            use_container_width=True
                        )
            else:
                st.error("❌ Не удалось сопоставить ни одного образца")
                
                # Показываем отладочную информацию
                st.markdown("### 🔍 Отладочная информация:")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Ключи из правильного порядка:**")
                    for sample in correct_samples[:10]:
                        st.write(f"- {sample['key']} → '{sample['correct_name']}'")
                
                with col2:
                    st.write("**Ключи из химического анализа:**")
                    all_analysis_samples = []
                    for samples in chemical_tables.values():
                        all_analysis_samples.extend(samples)
                    for sample in all_analysis_samples[:10]:
                        st.write(f"- {sample['key']} → '{sample['original_name']}'")
                
        except Exception as e:
            st.error(f"❌ Произошла ошибка при обработке: {str(e)}")
            st.info("💡 Убедитесь, что файлы имеют правильный формат и содержат таблицы в текстовом представлении")
    
    # Информация о проекте
    st.markdown("---")
    st.markdown("""
    ### ℹ️ О ПРОЕКТЕ
    
    Это веб-приложение автоматически сортирует результаты химического анализа 
    согласно заданному порядку образцов и обновляет названия на правильные.
    
    **Основные возможности:**
    - 📄 Поддержка форматов TXT и DOCX
    - 🔍 Автоматическое сопоставление образцов по цифрам и буквам в названиях
    - 👁️ Визуализация таблицы сопоставления для проверки
    - 📊 Обработка нескольких таблиц в одном файле
    - 📤 Экспорт результатов в Excel
    
    **Технологии:** Python, Streamlit, Pandas, OpenPyXL, python-docx
    """)

if __name__ == "__main__":
    main()
